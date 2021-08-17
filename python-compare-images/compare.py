# USAGE
# python compare.py

# import the necessary packages
from docx import Document
from docx.shared import Inches
from skimage.metrics import structural_similarity as ssim
import matplotlib.pyplot as plt
import numpy as np
import cv2
import os

document = Document()
document.add_heading('Report of testing', 0)
document.add_paragraph('MSE (Mean Squared Error) - value of 0 for MSE indicates perfect similarity')
document.add_paragraph('SSIM (Structural Similarity Index) - value of 1 indicates perfect similarity.')


def mse(imageA, imageB):
    # the 'Mean Squared Error' between the two images is the
    # sum of the squared difference between the two images;
    # NOTE: the two images must have the same dimension
    err = np.sum((imageA.astype("float") - imageB.astype("float")) ** 2)
    err /= float(imageA.shape[0] * imageA.shape[1])

    # return the MSE, the lower the error, the more "similar"
    # the two images are
    return err


def compare_images(imageA, imageB, title):
    # compute the mean squared error and structural similarity
    # index for the images
    m = mse(imageA, imageB)
    s = ssim(imageA, imageB)

    if m < 1400.00:
        print("MSE: %.2f, SSIM: %.2f" % (m, s))

        # setup the figure
        fig = plt.figure(title)
        plt.suptitle("MSE: %.2f, SSIM: %.2f" % (m, s))

        # show first image
        ax = fig.add_subplot(1, 2, 1)
        plt.imshow(imageA, cmap=plt.cm.gray)
        plt.axis("off")

        # show the second image
        ax = fig.add_subplot(1, 2, 2)
        plt.imshow(imageB)
        plt.axis("off")

        plt.savefig('foo.png')
        document.add_picture('foo.png', width=Inches(5))
        # show the images
        # plt.show()

    document.save('demo.docx')


listImages = os.listdir("images")

# load the images -- the original, the original + contrast,
# and the original + photoshop
original = cv2.imread("images/jp_gates_original.png")
contrast = cv2.imread("images/jp_gates_contrast.png")
shopped = cv2.imread("images/jp_gates_photoshopped.png")

# convert the images to grayscale
original = cv2.cvtColor(original, cv2.COLOR_BGR2GRAY)
contrast = cv2.cvtColor(contrast, cv2.COLOR_BGR2GRAY)
shopped = cv2.cvtColor(shopped, cv2.COLOR_BGR2GRAY)

# initialize the figure
fig = plt.figure("Images")
images = ("Original", original), ("Contrast", contrast), ("Photoshopped", shopped)

# loop over the images
for (i, (name, image)) in enumerate(images):
    # show the image
    ax = fig.add_subplot(1, 3, i + 1)
    ax.set_title(name)
    plt.imshow(image, cmap=plt.cm.gray)
    plt.axis("off")

# show the figure
# plt.show()


# compare the images
compare_images(original, original, "Original vs. Original")
compare_images(original, contrast, "Original vs. Contrast")
compare_images(original, shopped, "Original vs. Photoshopped")

if __name__ == '__main__':
    compare_images(original, original, "Original vs. Original")
    compare_images(original, contrast, "Original vs. Contrast")
    compare_images(original, shopped, "Original vs. Photoshopped")
